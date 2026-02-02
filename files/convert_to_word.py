
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_to_doc(doc, md_content, base_path="."):
    lines = md_content.split('\n')
    in_code_block = False
    code_content = []
    
    table_mode = False
    table_rows = []

    for line in lines:
        line = line.rstrip()
        
        # Code Blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                runner = p.add_run("\n".join(code_content))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                p.style = 'No Spacing'
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue

        # Tables (Basic detection: starts and ends with |)
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not table_mode:
                table_mode = True
                table_rows = []
            table_rows.append(line)
            continue
        elif table_mode:
            # Table ended
            process_table(doc, table_rows)
            table_mode = False
            table_rows = []
        
        # Headers
        if line.startswith("#"):
            level = len(line.split(' ')[0])
            text = line[level:].strip()
            # Clean comments
            text = re.sub(r'<!--.*?-->', '', text)
            doc.add_heading(text, level=min(level, 9))
            continue
            
    # Images - handle potential leading whitespace
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            # Resolve relative path
            full_img_path = os.path.join(base_path, img_path)
            if os.path.exists(full_img_path):
                try:
                    doc.add_picture(full_img_path, width=Inches(6.0))
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Image could not be loaded: {img_path}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            continue

        # Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            format_text(p, text)
            continue
            
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            format_text(p, text)
            continue
            
        # Normal Text (ignore separators)
        if line.strip() == "---" or line.strip() == "":
            continue
            
        p = doc.add_paragraph()
        format_text(p, line)

    # Process pending table if file ends with table
    if table_mode:
        process_table(doc, table_rows)

def format_text(paragraph, text):
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic *text*
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                     run = paragraph.add_run(subpart[1:-1])
                     run.italic = True
                else:
                    paragraph.add_run(subpart)

def process_table(doc, rows):
    if not rows:
        return
    # Remove alignment row if present (monitor |---|)
    clean_rows = [r for r in rows if not re.search(r'^\s*\|?\s*:?-+:?\s*\|', r)]
    
    if not clean_rows:
        return

    # Determine max columns
    max_cols = 0
    parsed_rows = []
    for row in clean_rows:
        cols = [c.strip() for c in row.strip().strip('|').split('|')]
        if len(cols) > max_cols:
            max_cols = len(cols)
        parsed_rows.append(cols)
        
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(parsed_rows):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                # Bold header
                if i == 0:
                    for run in row_cells[j].paragraphs[0].runs:
                        run.bold = True

def convert_file(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    doc = Document()
    base_path = os.path.dirname(md_path)
    add_markdown_to_doc(doc, content, base_path)
    try:
        doc.save(docx_path)
        print(f"Created: {docx_path}")
    except PermissionError:
        print(f"Error: Could not save {docx_path}. Is the file open?")

if __name__ == "__main__":
    convert_file("PROJECT_EXPLANATION.md", "Project_Explanation_v2.docx")
    convert_file("TECHNICAL_REPORT.md", "Technical_Report_v2.docx")
